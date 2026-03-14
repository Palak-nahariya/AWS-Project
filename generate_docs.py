"""
Script to generate PROJECT_DOCUMENTATION.docx for Cloud Bank Analytics
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ───────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    hs.font.name = 'Calibri'
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A56DB')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'EBF0FA')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # light grey background via shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    shading.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(shading)


# ═════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('☁️ Cloud Bank Analytics')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Project Documentation')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run(
    'A comprehensive cloud-based banking application featuring secure account management,\n'
    'financial transactions, role-based analytics dashboards, and real-time fraud detection\n'
    '— powered by AWS services.'
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('© 2026 Cloud Bank Analytics  •  Powered by AWS')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (static)
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)

toc_items = [
    '1.  Project Overview',
    '2.  Technology Stack',
    '3.  Project Structure',
    '4.  Architecture & Data Flow',
    '5.  Database Schema',
    '6.  Module-Level Documentation',
    '7.  API Routes & Endpoints',
    '8.  User Roles & Access Control',
    '9.  Fraud Detection System',
    '10. Security Features',
    '11. Local Development Mode',
    '12. Environment Configuration',
    '13. AWS Infrastructure Setup',
    '14. Installation & Running',
    '15. Production Deployment',
    '16. UI & Design',
    '17. Testing',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)

doc.add_paragraph(
    'Cloud Bank Analytics is a full-stack banking web application built with Python Flask '
    'that demonstrates the integration of multiple AWS services for a real-world financial '
    'application. It includes:'
)

add_bullet(' — registration, login, deposits, withdrawals, transfers, and transaction history.', 'Core banking operations')
add_bullet(' — real-time fraud monitoring, custom financial report generation, and regulatory compliance tracking.', 'Role-based analytics dashboards')
add_bullet(' — an algorithm that scores every transaction (0–100) and can automatically freeze accounts.', 'Fraud detection engine')
add_bullet(' — DynamoDB for data persistence, SNS for push notifications, EC2 for hosting, and IAM for access control.', 'AWS-native backend')
add_bullet(' — JSON-file-based storage that lets developers run and test the entire application without any AWS credentials.', 'Local development mode')

# ═════════════════════════════════════════════════════════════════════════
#  2. TECHNOLOGY STACK
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Technology Stack', level=1)

add_table(
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Backend Framework', 'Flask 3.0.0', 'Web server, routing, templating'],
        ['Database', 'AWS DynamoDB', 'NoSQL data persistence'],
        ['Notifications', 'AWS SNS', 'Email / SMS alerts for fraud and compliance'],
        ['Hosting', 'AWS EC2', 'Production server'],
        ['Authentication', 'bcrypt 4.1.2', 'Password hashing with 12 rounds'],
        ['Session Management', 'Flask sessions', 'HTTP-only, SameSite cookies'],
        ['Environment Config', 'python-dotenv 1.0.0', '.env file management'],
        ['HTTP Server', 'Werkzeug 3.0.1 / Gunicorn 21.2.0', 'WSGI servers (dev / prod)'],
        ['Cloud SDK', 'boto3 1.34.0', 'AWS service integration'],
        ['Frontend', 'Jinja2, CSS, Chart.js', 'UI rendering and charts'],
    ],
    col_widths=[1.8, 2.2, 2.8]
)

doc.add_heading('Dependencies (requirements.txt)', level=2)
add_code_block(
    'Flask==3.0.0\n'
    'boto3==1.34.0\n'
    'bcrypt==4.1.2\n'
    'python-dotenv==1.0.0\n'
    'Werkzeug==3.0.1\n'
    'gunicorn==21.2.0'
)

# ═════════════════════════════════════════════════════════════════════════
#  3. PROJECT STRUCTURE
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Project Structure', level=1)

structure = """\
AWS Project/
├── app.py                     # Flask application entry point
├── config.py                  # Centralized configuration (Config class)
├── requirements.txt           # Python dependencies
├── .env / .env.example        # Environment variables
│
├── models/                    # Data access layer (DynamoDB)
│   ├── __init__.py
│   ├── user.py                # User model – CRUD, auth, password hashing
│   ├── account.py             # Account model – balance, freeze/activate
│   └── transaction.py         # Transaction model – CRUD, fraud scoring
│
├── simple_models.py           # Lightweight wrappers for local-storage mode
├── local_storage.py           # JSON-file-based mock database
├── local_data.json            # Persisted local data file
│
├── services/                  # Business logic layer
│   ├── auth_service.py        # Registration, login, session management
│   ├── banking_service.py     # Deposit, withdraw, transfer, history
│   ├── analytics_service.py   # Fraud dashboard, reports, compliance
│   └── notification_service.py # AWS SNS alert delivery
│
├── routes/                    # HTTP route handlers (Flask Blueprints)
│   ├── auth_routes.py         # /, /register, /login, /logout
│   ├── account_routes.py      # /dashboard
│   ├── transaction_routes.py  # /deposit, /withdraw, /transfer, /history
│   └── analytics_routes.py    # /analytics/* (fraud, reports, compliance)
│
├── templates/                 # Jinja2 HTML templates
│   ├── base.html              # Master layout
│   ├── index / login / register / dashboard / deposit / withdraw / transfer / history .html
│   └── analytics/
│       ├── fraud_monitoring.html
│       ├── reports.html
│       └── compliance.html
│
├── static/css/style.css       # Global stylesheet (dark glassmorphism)
│
├── test_aws_connection.py     # AWS connectivity test
├── test_password_fix.py       # Password hashing test
├── debug_passwords.py         # Password debug utility
└── generate_hash.py           # Standalone hash generator"""

add_code_block(structure)

# ═════════════════════════════════════════════════════════════════════════
#  4. ARCHITECTURE & DATA FLOW
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Architecture & Data Flow', level=1)

doc.add_paragraph(
    'The application follows a clean three-layer architecture: Routes → Services → Models/Storage.'
)

arch = """\
CLIENT (Browser)
    │ HTTP requests
    ▼
FLASK APP (app.py)
    ├── auth_bp (routes)
    ├── account_bp (routes)
    ├── transaction_bp (routes)
    └── analytics_bp (routes)
         │
         ▼
    SERVICE LAYER
    ├── AuthService
    ├── BankingService
    ├── AnalyticsService
    └── NotificationService
         │
         ▼
    MODEL / DATA LAYER
    ├── User, Account, Transaction  ──►  AWS DynamoDB (3 tables)
    ├── LocalStorage (fallback)     ──►  local_data.json
    └── NotificationService         ──►  AWS SNS (3 topics)"""

add_code_block(arch)

doc.add_heading('Layer Responsibilities', level=2)
add_table(
    ['Layer', 'Files', 'Responsibility'],
    [
        ['Routes', 'routes/*.py', 'HTTP handling, input validation, template rendering'],
        ['Services', 'services/*.py', 'Business logic, orchestration, fraud decisions'],
        ['Models', 'models/*.py, simple_models.py', 'Data access, DynamoDB / local-storage abstraction'],
        ['Storage', 'local_storage.py, local_data.json', 'JSON-file mock database for local dev'],
    ]
)

# ═════════════════════════════════════════════════════════════════════════
#  5. DATABASE SCHEMA
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Database Schema', level=1)

doc.add_heading('5.1  Users Table (CloudBank_Users)', level=2)
add_table(
    ['Attribute', 'Type', 'Description'],
    [
        ['UserID', 'String (PK)', 'UUID, primary key'],
        ['Name', 'String', 'Full name'],
        ['Email', 'String', 'Unique email (GSI: EmailIndex)'],
        ['PasswordHash', 'String', 'bcrypt hash'],
        ['Role', 'String', 'customer / analyst / manager / compliance'],
        ['CreatedAt', 'String (ISO 8601)', 'Creation timestamp'],
        ['UpdatedAt', 'String (ISO 8601)', 'Last update timestamp'],
    ]
)
doc.add_paragraph('Global Secondary Index: EmailIndex (PK: Email) — used for login lookups.')

doc.add_heading('5.2  Accounts Table (CloudBank_Accounts)', level=2)
add_table(
    ['Attribute', 'Type', 'Description'],
    [
        ['AccountID', 'String (PK)', 'UUID'],
        ['UserID', 'String', 'Owner reference'],
        ['Balance', 'Number (Decimal)', 'Current balance'],
        ['AccountType', 'String', 'SAVINGS / CHECKING'],
        ['Status', 'String', 'ACTIVE / FROZEN'],
        ['CreatedAt', 'String', 'Timestamp'],
        ['UpdatedAt', 'String', 'Timestamp'],
    ]
)
doc.add_paragraph('Global Secondary Index: UserIDIndex (PK: UserID) — fetch all accounts for a user.')

doc.add_heading('5.3  Transactions Table (CloudBank_Transactions)', level=2)
add_table(
    ['Attribute', 'Type', 'Description'],
    [
        ['TransactionID', 'String (PK)', 'UUID'],
        ['Date', 'String (SK)', 'ISO 8601 timestamp (sort key)'],
        ['AccountID', 'String', 'Source account'],
        ['TargetAccountID', 'String', 'Destination (transfers only)'],
        ['TransactionType', 'String', 'DEPOSIT / WITHDRAW / TRANSFER'],
        ['Amount', 'Number (Decimal)', 'Transaction amount'],
        ['Status', 'String', 'PENDING / COMPLETED / FAILED'],
        ['FraudScore', 'Number', '0–100 risk score'],
        ['Description', 'String', 'User-provided note'],
    ]
)
doc.add_paragraph(
    'Global Secondary Indexes: AccountIDIndex (PK: AccountID, SK: Date), '
    'DateIndex (PK: Date), FraudScoreIndex (PK: FraudScore).'
)

# ═════════════════════════════════════════════════════════════════════════
#  6. MODULE-LEVEL DOCUMENTATION
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Module-Level Documentation', level=1)

# 6.1 app.py
doc.add_heading('6.1  app.py — Application Entry Point', level=2)
add_bullet('Creates the Flask app and loads configuration from Config.', '')
add_bullet('Registers four blueprints: auth_bp, account_bp, transaction_bp, analytics_bp.', '')
add_bullet('Configures session lifetime (default 30 min).', '')
add_bullet('Defines 404 and 500 error handlers.', '')
add_bullet('Runs on 0.0.0.0:5000 in debug mode.', '')

# 6.2 config.py
doc.add_heading('6.2  config.py — Configuration', level=2)
doc.add_paragraph('Loads environment variables from .env via python-dotenv and exposes them through the Config class.')
add_table(
    ['Config Group', 'Key Settings'],
    [
        ['Flask', 'SECRET_KEY, DEBUG'],
        ['Storage mode', 'USE_LOCAL_STORAGE (default True)'],
        ['AWS', 'AWS_REGION, access keys'],
        ['DynamoDB tables', 'CloudBank_Users, CloudBank_Accounts, CloudBank_Transactions'],
        ['SNS ARNs', 'Transaction, Compliance, and System alert topics'],
        ['Session', '30-min lifetime, HTTP-only cookies, SameSite=Lax'],
        ['Security', 'BCRYPT_ROUNDS = 12'],
        ['Fraud thresholds', 'Alert ≥ 70, Freeze ≥ 90'],
    ]
)

# 6.3 User Model
doc.add_heading('6.3  models/user.py — User Model', level=2)
add_table(
    ['Method', 'Description'],
    [
        ['create_user()', 'Hashes password with bcrypt, stores in DynamoDB or local storage'],
        ['get_user_by_id()', 'Fetch user by primary key'],
        ['get_user_by_email()', 'Fetch via EmailIndex GSI'],
        ['verify_password()', 'Compare plaintext against bcrypt hash'],
        ['authenticate()', 'Full login flow: lookup → verify → return sanitized user'],
        ['update_user()', 'Partial update (prevents overwriting UserID and PasswordHash)'],
    ]
)

# 6.4 Account Model
doc.add_heading('6.4  models/account.py — Account Model', level=2)
add_table(
    ['Method', 'Description'],
    [
        ['create_account()', 'New account with conditional check for duplicates'],
        ['get_account()', 'Fetch by AccountID'],
        ['get_accounts_by_user()', 'Query UserIDIndex GSI'],
        ['update_balance()', 'Atomic ADD / SUBTRACT with sufficient-balance check'],
        ['freeze_account()', 'Set status to FROZEN'],
        ['activate_account()', 'Restore status to ACTIVE'],
    ]
)

# 6.5 Transaction Model
doc.add_heading('6.5  models/transaction.py — Transaction Model', level=2)
add_table(
    ['Method', 'Description'],
    [
        ['create_transaction()', 'Records transaction with auto-calculated fraud score'],
        ['update_transaction_status()', 'Mark as COMPLETED or FAILED'],
        ['get_transaction()', 'Fetch single transaction'],
        ['get_account_transactions()', 'History for an account (most recent first)'],
        ['get_high_fraud_transactions()', 'Query by fraud score threshold'],
        ['get_transactions_by_date_range()', 'Date-range query via DateIndex'],
        ['_calculate_fraud_score()', 'Internal scoring algorithm (see §9)'],
    ]
)

# 6.6 AuthService
doc.add_heading('6.6  services/auth_service.py — AuthService', level=2)
add_table(
    ['Method', 'Description'],
    [
        ['register()', 'Validates input → checks duplicate email → creates user'],
        ['login()', 'Validates input → authenticates → returns user data'],
        ['get_user()', 'Session re-validation (strips PasswordHash)'],
        ['update_profile()', 'Filters out sensitive fields before update'],
    ]
)

# 6.7 BankingService
doc.add_heading('6.7  services/banking_service.py — BankingService', level=2)
doc.add_paragraph('Automatically selects simple_models (local) or models (DynamoDB) based on USE_LOCAL_STORAGE.')
add_table(
    ['Method', 'Description'],
    [
        ['create_account()', 'Creates account with UUID'],
        ['get_user_accounts()', 'Returns all accounts for a user'],
        ['deposit()', 'Creates transaction → updates balance → marks complete'],
        ['withdraw()', 'Creates transaction → checks fraud → freezes if ≥ 90 → updates balance'],
        ['transfer()', 'Validates target → deducts source → credits target → rollback on failure'],
        ['get_transaction_history()', 'Delegates to transaction model'],
    ]
)

# 6.8 AnalyticsService
doc.add_heading('6.8  services/analytics_service.py — AnalyticsService', level=2)

doc.add_paragraph('Scenario 1 — Fraud Monitoring (Analyst):')
add_table(
    ['Method', 'Description'],
    [
        ['get_fraud_monitoring_dashboard()', 'Categorizes flagged txns: Critical (≥90), High (80–89), Medium (70–79)'],
        ['get_recent_transactions_feed()', 'Live feed of last N hours'],
        ['investigate_transaction()', 'Deep dive: transaction + account + recent history'],
    ]
)

doc.add_paragraph('Scenario 2 — Reports (Manager):')
add_table(
    ['Method', 'Description'],
    [
        ['generate_financial_report()', 'Aggregate metrics: deposit/withdrawal/transfer volumes'],
        ['get_deposit_growth_trends()', 'Daily deposit totals & averages over N days'],
        ['get_transaction_volume_analysis()', 'Hourly volume distribution, peak hour detection'],
    ]
)

doc.add_paragraph('Scenario 3 — Compliance (Officer):')
add_table(
    ['Method', 'Description'],
    [
        ['get_compliance_dashboard()', 'Monitors large txns (>$10K), suspicious activity, failure rate'],
        ['drill_down_compliance_metric()', 'Filtered transaction lists per metric type'],
    ]
)

# 6.9 NotificationService
doc.add_heading('6.9  services/notification_service.py — NotificationService', level=2)
add_table(
    ['Method', 'Description'],
    [
        ['send_transaction_alert()', 'Publishes to Transaction Alerts SNS topic'],
        ['send_compliance_alert()', 'Publishes to Compliance Alerts SNS topic'],
        ['send_system_alert()', 'Publishes to System Alerts SNS topic'],
        ['notify_high_fraud_transaction()', 'Formats and sends fraud alert with details'],
        ['notify_account_frozen()', 'Sends account-freeze notification'],
    ]
)

# 6.10 LocalStorage
doc.add_heading('6.10  local_storage.py — LocalStorage', level=2)
doc.add_paragraph(
    'A JSON-file-based mock database (local_data.json) that mirrors the DynamoDB API surface. '
    'Used when USE_LOCAL_STORAGE=True. In-memory dict backed by local_data.json. '
    'Implements all CRUD for users, accounts, and transactions. Auto-saves on every write.'
)

# 6.11 simple_models
doc.add_heading('6.11  simple_models.py — SimpleAccount / SimpleTransaction', level=2)
doc.add_paragraph(
    'Lightweight wrappers that delegate to LocalStorage while matching the Account / Transaction '
    'model interface. Includes a simplified fraud-score calculator (amount-based only).'
)

# ═════════════════════════════════════════════════════════════════════════
#  7. API ROUTES & ENDPOINTS
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('7. API Routes & Endpoints', level=1)

doc.add_heading('7.1  Authentication Routes', level=2)
add_table(
    ['Method', 'URL', 'Auth', 'Description'],
    [
        ['GET', '/', 'No', 'Landing page (redirects to dashboard if logged in)'],
        ['GET/POST', '/register', 'No', 'User registration form & handler'],
        ['GET/POST', '/login', 'No', 'Login form & handler'],
        ['GET', '/logout', 'Yes', 'Clears session, redirects to login'],
    ]
)

doc.add_heading('7.2  Account Routes', level=2)
add_table(
    ['Method', 'URL', 'Auth', 'Description'],
    [
        ['GET', '/dashboard', 'Yes', 'Shows all user accounts with balances'],
    ]
)

doc.add_heading('7.3  Transaction Routes', level=2)
add_table(
    ['Method', 'URL', 'Auth', 'Description'],
    [
        ['GET/POST', '/deposit', 'Yes', 'Deposit money into selected account'],
        ['GET/POST', '/withdraw', 'Yes', 'Withdraw money (fraud check + SNS alert)'],
        ['GET/POST', '/transfer', 'Yes', 'Transfer between accounts (fraud + rollback)'],
        ['GET', '/history', 'Yes', 'View transaction history with account selector'],
    ]
)

doc.add_heading('7.4  Analytics Routes (prefix: /analytics)', level=2)
add_table(
    ['Method', 'URL', 'Role', 'Description'],
    [
        ['GET', '/analytics/fraud-monitoring', 'Analyst', 'Fraud monitoring dashboard'],
        ['GET', '/analytics/api/recent-transactions', 'Analyst', 'JSON: live transaction feed'],
        ['GET', '/analytics/api/investigate/<id>', 'Analyst', 'JSON: transaction deep-dive'],
        ['POST', '/analytics/api/approve-transaction', 'Analyst', 'Activate a frozen account'],
        ['POST', '/analytics/api/freeze-account', 'Analyst', 'Freeze a suspicious account'],
        ['GET', '/analytics/reports', 'Manager', 'Report generation dashboard'],
        ['POST', '/analytics/api/financial-report', 'Manager', 'JSON: financial metrics'],
        ['GET', '/analytics/api/deposit-trends', 'Manager', 'JSON: deposit growth trends'],
        ['GET', '/analytics/api/transaction-volume', 'Manager', 'JSON: hourly volume distribution'],
        ['GET', '/analytics/compliance', 'Compliance', 'Compliance monitoring dashboard'],
        ['GET', '/analytics/api/compliance-drilldown', 'Compliance', 'JSON: compliance detail'],
    ]
)

# ═════════════════════════════════════════════════════════════════════════
#  8. USER ROLES & ACCESS CONTROL
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('8. User Roles & Access Control', level=1)

add_table(
    ['Role', 'Access'],
    [
        ['Customer', 'Dashboard, Deposit, Withdraw, Transfer, History'],
        ['Analyst', 'All Customer pages + Fraud Monitoring (/analytics/fraud-monitoring)'],
        ['Manager', 'All Customer pages + Custom Reports (/analytics/reports)'],
        ['Compliance', 'All Customer pages + Regulatory Compliance (/analytics/compliance)'],
    ]
)

doc.add_heading('Access Control Implementation', level=2)
add_bullet('login_required decorator: checks session["user_id"].', '')
add_bullet('role_required(role) decorator: checks session["role"]; admin role bypasses all checks.', '')
add_bullet('Navbar dynamically shows role-specific links based on session.role.', '')

# ═════════════════════════════════════════════════════════════════════════
#  9. FRAUD DETECTION SYSTEM
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Fraud Detection System', level=1)

doc.add_heading('9.1  Scoring Algorithm', level=2)
doc.add_paragraph('Every withdrawal and transfer is scored on a 0–100 scale:')
add_table(
    ['Factor', 'Condition', 'Points'],
    [
        ['Amount Anomaly (max 40)', '> 3× user average', '40'],
        ['', '> 2× user average', '25'],
        ['', '> 1.5× user average', '15'],
        ['Frequency (max 30)', '> 20 txns in 24h', '30'],
        ['', '> 10 txns in 24h', '20'],
        ['', '> 5 txns in 24h', '10'],
        ['Large Amount (max 30)', '> $10,000', '30'],
        ['', '> $5,000', '20'],
        ['', '> $2,000', '10'],
    ]
)

doc.add_heading('9.2  Automated Actions', level=2)
add_table(
    ['Threshold', 'Action'],
    [
        ['Score ≥ 70', 'SNS alert sent to Transaction Alerts topic'],
        ['Score ≥ 90', 'Account automatically frozen, transaction blocked'],
    ]
)

doc.add_heading('9.3  Analyst Dashboard', level=2)
doc.add_paragraph('Analysts review flagged transactions categorized as:')
add_bullet('Critical (score ≥ 90) — account frozen', '🔴 ')
add_bullet('High (80–89) — needs investigation', '🟠 ')
add_bullet('Medium (70–79) — flagged for review', '🟡 ')
doc.add_paragraph('Analysts can freeze or approve/activate accounts via API endpoints.')

# ═════════════════════════════════════════════════════════════════════════
#  10. SECURITY FEATURES
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Security Features', level=1)

add_table(
    ['Feature', 'Implementation'],
    [
        ['Password Hashing', 'bcrypt with 12 rounds (configurable via BCRYPT_ROUNDS)'],
        ['Session Security', 'HTTP-only cookies, SameSite=Lax, 30-min expiry'],
        ['Input Validation', 'Server-side validation in services layer'],
        ['Sensitive Data Protection', 'PasswordHash stripped from all API responses'],
        ['Fraud Prevention', 'Real-time scoring, automatic account freezing'],
        ['SNS Alerts', 'Immediate notification for high-risk transactions'],
        ['Account Freeze', 'Frozen accounts cannot process transactions'],
        ['Transfer Rollback', 'Automatic balance restoration if target credit fails'],
        ['Conditional Writes', 'DynamoDB condition expressions prevent race conditions'],
    ]
)

# ═════════════════════════════════════════════════════════════════════════
#  11. LOCAL DEVELOPMENT MODE
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Local Development Mode', level=1)

doc.add_paragraph('When USE_LOCAL_STORAGE=True (default), the application runs entirely without AWS:')

add_table(
    ['Component', 'Production', 'Local Mode'],
    [
        ['User data', 'DynamoDB CloudBank_Users', 'local_data.json'],
        ['Account data', 'DynamoDB CloudBank_Accounts', 'local_data.json'],
        ['Transaction data', 'DynamoDB CloudBank_Transactions', 'local_data.json'],
        ['Models used', 'models/ (full DynamoDB)', 'simple_models.py + local_storage.py'],
        ['Fraud scoring', 'Full 3-factor algorithm', 'Simplified amount-only check'],
        ['Notifications', 'SNS topics', 'Console print (no-op)'],
    ]
)

doc.add_heading('How It Works', level=2)
add_bullet('config.py reads USE_LOCAL_STORAGE from .env.', '1. ')
add_bullet('Each model class checks this flag in __init__().', '2. ')
add_bullet('If local mode, it imports local_storage.local_db (singleton instance).', '3. ')
add_bullet('banking_service.py switches between simple_models and models at import time.', '4. ')
add_bullet('All data persists to local_data.json via json.dump().', '5. ')

# ═════════════════════════════════════════════════════════════════════════
#  12. ENVIRONMENT CONFIGURATION
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('12. Environment Configuration', level=1)

doc.add_paragraph('The .env file controls all application settings:')

env_block = """\
# Storage Mode
USE_LOCAL_STORAGE=True

# AWS Credentials
AWS_REGION=us-east-1
AWS_ACCESS_KEY_ID=your_access_key
AWS_SECRET_ACCESS_KEY=your_secret_key

# DynamoDB Tables
DYNAMODB_USERS_TABLE=CloudBank_Users
DYNAMODB_ACCOUNTS_TABLE=CloudBank_Accounts
DYNAMODB_TRANSACTIONS_TABLE=CloudBank_Transactions

# SNS Topics
SNS_TRANSACTION_ALERTS_ARN=arn:aws:sns:...
SNS_COMPLIANCE_ALERTS_ARN=arn:aws:sns:...
SNS_SYSTEM_ALERTS_ARN=arn:aws:sns:...

# Flask
FLASK_SECRET_KEY=your_secret_key_here
DEBUG=True

# Security
SESSION_TIMEOUT=1800
BCRYPT_ROUNDS=12

# Fraud Detection
FRAUD_ALERT_THRESHOLD=70
FRAUD_FREEZE_THRESHOLD=90"""
add_code_block(env_block)

# ═════════════════════════════════════════════════════════════════════════
#  13. AWS INFRASTRUCTURE SETUP
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('13. AWS Infrastructure Setup', level=1)

doc.add_heading('Step 1 — Create DynamoDB Tables', level=2)

doc.add_paragraph('Users Table:')
add_code_block('Table Name: CloudBank_Users\nPrimary Key: UserID (String)\nGSI: EmailIndex → Email (PK)')

doc.add_paragraph('Accounts Table:')
add_code_block('Table Name: CloudBank_Accounts\nPrimary Key: AccountID (String)\nGSI: UserIDIndex → UserID (PK)')

doc.add_paragraph('Transactions Table:')
add_code_block(
    'Table Name: CloudBank_Transactions\n'
    'Primary Key: TransactionID (String)\n'
    'Sort Key: Date (String)\n'
    'GSI: AccountIDIndex → AccountID (PK), Date (SK)\n'
    'GSI: DateIndex → Date (PK)\n'
    'GSI: FraudScoreIndex → FraudScore (Number PK)'
)

doc.add_heading('Step 2 — Create SNS Topics', level=2)
add_table(
    ['Topic Name', 'Purpose'],
    [
        ['TransactionAlerts', 'Fraud detection notifications'],
        ['ComplianceAlerts', 'Compliance threshold warnings'],
        ['SystemAlerts', 'System health monitoring'],
    ]
)
doc.add_paragraph('Subscribe email addresses to each topic for notifications.')

doc.add_heading('Step 3 — Configure IAM', level=2)
doc.add_paragraph('Create an IAM role/user with these policies:')
add_bullet('AmazonDynamoDBFullAccess', '')
add_bullet('AmazonSNSFullAccess', '')
doc.add_paragraph('Attach to the EC2 instance or generate access keys for local development.')

doc.add_heading('Step 4 — Update .env', level=2)
doc.add_paragraph('Set USE_LOCAL_STORAGE=False and fill in all AWS credentials and ARNs.')

# ═════════════════════════════════════════════════════════════════════════
#  14. INSTALLATION & RUNNING
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('14. Installation & Running', level=1)

doc.add_heading('Prerequisites', level=2)
add_bullet('Python 3.8 or later', '')
add_bullet('pip package manager', '')
add_bullet('(Optional) AWS account with configured credentials', '')

doc.add_heading('Setup Steps', level=2)

setup_block = """\
# 1. Navigate to project directory
cd "C:\\Users\\khurai\\OneDrive\\Desktop\\AWS Project"

# 2. Create virtual environment
python -m venv venv

# 3. Activate virtual environment (Windows)
venv\\Scripts\\activate

# 4. Install dependencies
pip install -r requirements.txt

# 5. Configure environment
copy .env.example .env
# Edit .env with your settings

# 6. Run the application
python app.py"""
add_code_block(setup_block)

doc.add_paragraph('Access the application at http://localhost:5000')

doc.add_heading('First-Time Registration', level=2)
add_bullet('Go to /register.', '1. ')
add_bullet('Fill in name, email, password (min 8 chars), and select a role.', '2. ')
add_bullet('A bank account with $1,000 initial balance is auto-created.', '3. ')
add_bullet('Log in at /login.', '4. ')

# ═════════════════════════════════════════════════════════════════════════
#  15. PRODUCTION DEPLOYMENT
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('15. Production Deployment', level=1)

doc.add_heading('EC2 Deployment Steps', level=2)
add_bullet('Launch an EC2 instance (t2.small or larger).', '1. ')
add_bullet('Install Python 3.8+ and project dependencies.', '2. ')
add_bullet('Set USE_LOCAL_STORAGE=False and DEBUG=False in .env.', '3. ')
add_bullet('Set a strong FLASK_SECRET_KEY.', '4. ')
add_bullet('Run via Gunicorn:', '5. ')
add_code_block('gunicorn -w 4 -b 0.0.0.0:5000 app:app')
add_bullet('Configure Nginx as a reverse proxy.', '6. ')
add_bullet('Enable HTTPS with Let\'s Encrypt / ACM.', '7. ')
add_bullet('Set up a systemd service for auto-restart.', '8. ')

doc.add_heading('Production Security Checklist', level=2)
checklist = [
    'HTTPS enabled',
    'SESSION_COOKIE_SECURE = True',
    'Strong FLASK_SECRET_KEY',
    'DynamoDB encryption at rest',
    'Restricted IAM permissions (least-privilege)',
    'Security groups: only ports 80/443 open',
    'CloudWatch logging enabled',
    'CSRF protection',
]
for item in checklist:
    add_bullet(item, '☐ ')

# ═════════════════════════════════════════════════════════════════════════
#  16. UI & DESIGN
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('16. UI & Design', level=1)

add_bullet(' Modern dark glassmorphism with semi-transparent cards and gradient backgrounds.', 'Theme:')
add_bullet(' Inter (Google Fonts) — weights 300–700.', 'Typography:')
add_bullet(' CSS fade-in effects and hover transitions.', 'Animations:')
add_bullet(' Chart.js for analytics dashboards.', 'Charts:')
add_bullet(' Mobile-friendly layout.', 'Responsive:')
add_bullet(' Color-coded flash messages (success, warning, danger, info).', 'Alerts:')
add_bullet(' Dynamic role-based links.', 'Navbar:')

# ═════════════════════════════════════════════════════════════════════════
#  17. TESTING
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('17. Testing', level=1)

doc.add_heading('Test Files', level=2)
add_table(
    ['File', 'Purpose'],
    [
        ['test_aws_connection.py', 'Verifies AWS DynamoDB and SNS connectivity'],
        ['test_password_fix.py', 'Tests bcrypt password hashing and verification'],
        ['debug_passwords.py', 'Inspects stored password hashes for debugging'],
        ['generate_hash.py', 'Standalone bcrypt hash generator utility'],
    ]
)

doc.add_heading('Running Tests', level=2)
add_code_block('python test_aws_connection.py    # Test AWS connectivity\npython test_password_fix.py      # Verify password hashing')

# ── Footer ───────────────────────────────────────────────────────────────
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('© 2026 Cloud Bank Analytics — Powered by AWS')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ═════════════════════════════════════════════════════════════════════════
#  SAVE
# ═════════════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'PROJECT_DOCUMENTATION.docx')
doc.save(output_path)
print(f'✅ Word document generated: {output_path}')
